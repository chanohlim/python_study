import PyPDF2,os, time

os.chdir('/Users/chanlim/Desktop/ /code/python_practice/pdf_word_python')
'''
pdfReader = PyPDF2.PdfReader('meetingminutes.pdf')
print(len(pdfReader.pages))
page1 = pdfReader.pages[0]
print(page1.extract_text(0))
'''

'''
pdf = PyPDF2.PdfReader('encrypted.pdf')
print(pdf.is_encrypted)

pdf.decrypt('rosebud')

page1 = pdf.pages[0]
print(page1.extract_text(0))
'''
'''
pdf1 = PyPDF2.PdfReader('meetingminutes.pdf')
pdf2 = PyPDF2.PdfReader('meetingminutes2.pdf')

pdfWriter = PyPDF2.PdfWriter()

for pg in range(len(pdf1.pages)):
    po = pdf1.pages[pg]
    pdfWriter.add_page(po)

for pg in range(len(pdf2.pages)):
    po = pdf2.pages[pg]
    pdfWriter.add_page(po)


pdfWriter.write("merged.pdf")
pdfWriter.close()
'''

''' # merge 2
merger = PyPDF2.PdfWriter()

for pdf in ["meetingminutes.pdf", "meetingminutes2.pdf"]:
    merger.append(pdf)

merger.write("merged2.pdf")
merger.close()
'''

'''
pdfWriter = PyPDF2.PdfWriter()
pdfWriter.append('meetingminutes.pdf')

pdfWriter.encrypt('helloworld')
pdfWriter.write('encrypted_meetingminutes.pdf')
pdfWriter.close()
'''

import docx

'''
def getText(filename):
    doc = docx.Document(filename)
    fullText = list()

    for p in doc.paragraphs:
        fullText.append(p.text)

    return '\n'.join(fullText)
'''

'''
doc = docx.Document('demo.docx')
doc.paragraphs[0].style = 'Normal'
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.add_paragraph('Hello, World!')
doc.save('restyled.docx')
'''

'''
doc = docx.Document()
doc.add_paragraph("hello")
paraObj1 = doc.add_paragraph('hi')
paraObj2 = doc.add_paragraph('what is up')
paraObj1.add_run('This text is added')

for i in range(1000):
    doc.add_paragraph('hi')

doc.save('multipleparagraph.docx')
'''

'''//Brute-force
f = open('dictionary.txt','r')
word_dict = list()

for word in f.readlines():
    word_dict.append(word.strip('\n'))
    word_dict.append(word.lower().strip('\n'))

print(word_dict)

pdf = PyPDF2.PdfReader('encrypted.pdf')
print(pdf.is_encrypted)

start = time.time()
for word in word_dict:
    if pdf.decrypt(word):
        end = time.time()
        print(f'The decryption key of this file is {word}, it took {end - start} seconds to find the key!')
        print(pdf.pages[0].extract_text(0))
        break
    else:
        print(f'The word {word} is not the decryption key.')
'''

''' //과제
insert_pdf = PyPDF2.PdfReader('[AI_Network_Lab_인사이트(2022-14호)]_최신_AI_불확실성_정량화_동향_및_시사점.pdf')
original = PyPDF2.PdfReader('노인기준연령 상향조정 논의에 대한 비판과 대안.pdf')

print(original.pages[0].extract_text())



pdfWriter = PyPDF2.PdfWriter()

pg = int(input(f'삽입을 원하는 페이지를 지정 (최고 페이지:{len(original.pages)}): '))
search = input('원하는 키워드 입력: ')

cnt = 0

for i in range(pg):
    po = original.pages[i]
    pdfWriter.add_page(po)

for po in insert_pdf.pages:
    if search in po.extract_text():
        pdfWriter.add_page(po)
        cnt += 1
        
print(f'{cnt} page added')

for i in range(pg,len(original.pages)):
    po = original.pages[i]
    pdfWriter.add_page(po)


pdfWriter.write("성균관.pdf")
pdfWriter.close()


1. pdf 파일 2개
2. 조작용 pdf와 삽입용 pdf
3. 삽입할 페이지 지정
4. 특정 텍스트를 가진 페이지만 삽입
'''


'''
pdfFiles = list()
for f in os.listdir('.'):
    if f.endswith('.pdf'):
        pdfFiles.append(f)
pdfFiles.sort()
print(pdfFiles)

pdfW = PyPDF2.PdfWriter()

for f in pdfFiles:
    pdf = PyPDF2.PdfReader(f)

    if pdf.is_encrypted == True:
        if pdf.decrypt('rosebud'):
            continue
        else:
            pdf.decrypt('helloworld')
    
    for i in range(1,len(pdf.pages)):
        pdfW.add_page(pdf.pages[i])

pdfW.write('merged_newfile.pdf')
pdfW.close()
'''